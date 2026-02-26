"""
Report Generator Module
=======================
Generates professional Excel reports with screenshots and Word documents.
This module imports data from data_source.py and handles all report generation.
"""

import os
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import data generation function
from data_source import get_report_data


# Output file paths
OUTPUT_DIR = Path(__file__).parent / "output"
EXCEL_FILE = OUTPUT_DIR / "Enterprise_Report.xlsx"
SUMMARY_IMAGE = OUTPUT_DIR / "summary.png"
ORDERS_IMAGE = OUTPUT_DIR / "orders.png"
CHART_TOP_NDC_SPEND = OUTPUT_DIR / "chart_top_ndc_spend.png"
CHART_SAVINGS_BY_MONTH = OUTPUT_DIR / "chart_savings_by_month.png"
CHART_TOP_NDC_PIE = OUTPUT_DIR / "chart_top_ndc_pie.png"
WORD_FILE = OUTPUT_DIR / "Final_Report.docx"


def ensure_output_directory() -> None:
    """Create output directory if it doesn't exist."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def create_excel_report(summary_df: pd.DataFrame, orders_df: pd.DataFrame) -> str:
    """
    Create a professionally formatted Excel report.
    
    Args:
        summary_df: Summary data DataFrame.
        orders_df: Orders data DataFrame.
        
    Returns:
        Path to the created Excel file.
    """
    excel_path = str(EXCEL_FILE)
    
    with pd.ExcelWriter(excel_path, engine="xlsxwriter") as writer:
        workbook = writer.book
        
        # Define formats
        header_format = workbook.add_format({
            "bold": True,
            "bg_color": "#217346",  # Dark green (Excel green)
            "font_color": "#FFFFFF",  # White text
            "border": 1,
            "text_wrap": True,
            "valign": "vcenter",
            "align": "center",
        })
        
        cell_format = workbook.add_format({
            "border": 1,
            "valign": "vcenter",
        })
        
        currency_format = workbook.add_format({
            "border": 1,
            "valign": "vcenter",
            "num_format": "$#,##0.00",
        })
        
        number_format = workbook.add_format({
            "border": 1,
            "valign": "vcenter",
            "num_format": "#,##0",
        })
        
        # Green format for "Yes" invoiced
        green_format = workbook.add_format({
            "border": 1,
            "valign": "vcenter",
            "bg_color": "#C6EFCE",
            "font_color": "#006100",
        })
        
        # Red format for "No" invoiced
        red_format = workbook.add_format({
            "border": 1,
            "valign": "vcenter",
            "bg_color": "#FFC7CE",
            "font_color": "#9C0006",
        })
        
        # Write Summary sheet
        _write_sheet(
            writer=writer,
            df=summary_df,
            sheet_name="Summary",
            header_format=header_format,
            cell_format=cell_format,
            currency_format=currency_format,
            number_format=number_format,
            currency_cols=["Total Spend"],
            number_cols=["Total Orders"],
        )
        
        # Write Orders sheet
        _write_sheet(
            writer=writer,
            df=orders_df,
            sheet_name="Orders",
            header_format=header_format,
            cell_format=cell_format,
            currency_format=currency_format,
            number_format=number_format,
            currency_cols=["Price", "Order Value"],
            number_cols=["Units"],
        )
        
        # Set green tab colors for both sheets
        writer.sheets["Summary"].set_tab_color("#00B050")
        writer.sheets["Orders"].set_tab_color("#00B050")
    
    print(f"✓ Excel report created: {excel_path}")
    return excel_path


def _write_sheet(
    writer: pd.ExcelWriter,
    df: pd.DataFrame,
    sheet_name: str,
    header_format,
    cell_format,
    currency_format,
    number_format,
    currency_cols: list = None,
    number_cols: list = None,
    conditional_col: str = None,
    green_format=None,
    red_format=None,
) -> None:
    """
    Write a DataFrame to an Excel sheet with professional formatting.
    
    Args:
        writer: ExcelWriter object.
        df: DataFrame to write.
        sheet_name: Name of the sheet.
        header_format: Format for header row.
        cell_format: Default cell format.
        currency_format: Format for currency columns.
        number_format: Format for number columns.
        currency_cols: List of column names to format as currency.
        number_cols: List of column names to format as numbers.
        conditional_col: Column name for conditional formatting.
        green_format: Format for "Yes" values.
        red_format: Format for "No" values.
    """
    currency_cols = currency_cols or []
    number_cols = number_cols or []
    
    # Write data starting at row 1 (row 0 for headers)
    df.to_excel(writer, sheet_name=sheet_name, index=False, startrow=1, header=False)
    
    worksheet = writer.sheets[sheet_name]
    
    # Write headers with formatting
    for col_num, column_name in enumerate(df.columns):
        worksheet.write(0, col_num, column_name, header_format)
    
    # Apply cell formatting
    for row_num in range(len(df)):
        for col_num, column_name in enumerate(df.columns):
            value = df.iloc[row_num, col_num]
            
            # Determine format
            if column_name in currency_cols:
                fmt = currency_format
            elif column_name in number_cols:
                fmt = number_format
            elif conditional_col and column_name == conditional_col:
                if value == "Yes":
                    fmt = green_format
                elif value == "No":
                    fmt = red_format
                else:
                    fmt = cell_format
            else:
                fmt = cell_format
            
            worksheet.write(row_num + 1, col_num, value, fmt)
    
    # Set column widths based on content
    for col_num, column_name in enumerate(df.columns):
        # Calculate max width
        max_length = max(
            len(str(column_name)),
            df[column_name].astype(str).str.len().max()
        )
        # Add padding and set width (max 50)
        worksheet.set_column(col_num, col_num, min(max_length + 3, 50))
    
    # Freeze header row
    worksheet.freeze_panes(1, 0)
    
    # Enable autofilter
    worksheet.autofilter(0, 0, len(df), len(df.columns) - 1)


def create_excel_screenshots(excel_path: str) -> tuple:
    """
    Convert Excel sheets to high-resolution PNG images.
    
    Args:
        excel_path: Path to the Excel file.
        
    Returns:
        Tuple of (summary_image_path, orders_image_path).
    """
    try:
        import excel2img
        
        summary_path = str(SUMMARY_IMAGE)
        orders_path = str(ORDERS_IMAGE)
        
        # Export Summary sheet
        excel2img.export_img(
            excel_path,
            summary_path,
            "Summary",
            None  # Export entire sheet
        )
        print(f"✓ Summary screenshot created: {summary_path}")
        
        # Export Orders sheet
        excel2img.export_img(
            excel_path,
            orders_path,
            "Orders",
            None
        )
        print(f"✓ Orders screenshot created: {orders_path}")
        
        return summary_path, orders_path
        
    except ImportError:
        print("⚠ excel2img not available. Using alternative screenshot method...")
        return _create_dataframe_images()
    except Exception as e:
        print(f"⚠ excel2img failed ({e}). Using alternative screenshot method...")
        return _create_dataframe_images()


def _create_dataframe_images() -> tuple:
    """
    Alternative method to create table images using matplotlib.
    Used as fallback when excel2img is not available.
    
    Returns:
        Tuple of (summary_image_path, orders_image_path).
    """
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # Use non-interactive backend
    
    # Get fresh data
    summary_df, orders_df = get_report_data()
    
    summary_path = str(SUMMARY_IMAGE)
    orders_path = str(ORDERS_IMAGE)
    
    # Create Summary image
    _create_table_image(
        df=summary_df,
        output_path=summary_path,
        title="Summary",
        currency_cols=["Total Spend"]
    )
    print(f"✓ Summary screenshot created: {summary_path}")
    
    # Create Orders image (limit rows for readability)
    _create_table_image(
        df=orders_df.head(25),  # Show first 25 rows
        output_path=orders_path,
        title="Orders",
        currency_cols=["Price", "Order Value"]
    )
    print(f"✓ Orders screenshot created: {orders_path}")
    
    return summary_path, orders_path


def _create_table_image(
    df: pd.DataFrame,
    output_path: str,
    title: str,
    currency_cols: list = None,
    highlight_col: str = None
) -> None:
    """
    Create a professional-looking table image from a DataFrame.
    
    Args:
        df: DataFrame to render.
        output_path: Path to save the image.
        title: Title for the table.
        currency_cols: Columns to format as currency.
        highlight_col: Column for conditional coloring.
    """
    import matplotlib.pyplot as plt
    import numpy as np
    
    currency_cols = currency_cols or []
    
    # Format currency columns
    df_display = df.copy()
    for col in currency_cols:
        if col in df_display.columns:
            df_display[col] = df_display[col].apply(lambda x: f"${x:,.2f}")
    
    # Calculate figure size based on data
    n_rows, n_cols = df_display.shape
    fig_width = max(12, n_cols * 2)
    fig_height = max(4, (n_rows + 2) * 0.4)
    
    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.axis('off')
    
    # Create table
    table = ax.table(
        cellText=df_display.values,
        colLabels=df_display.columns,
        cellLoc='center',
        loc='center',
        colColours=['#217346'] * n_cols  # Dark green headers
    )
    
    # Style the table
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1.2, 1.5)
    
    # Style header row
    for j in range(n_cols):
        cell = table[(0, j)]
        cell.set_text_props(weight='bold', color='white')
        cell.set_facecolor('#217346')
    
    # Apply conditional formatting for highlight column
    if highlight_col and highlight_col in df.columns:
        col_idx = df.columns.get_loc(highlight_col)
        for i in range(n_rows):
            cell = table[(i + 1, col_idx)]
            value = df.iloc[i][highlight_col]
            if value == "Yes":
                cell.set_facecolor('#C6EFCE')
            elif value == "No":
                cell.set_facecolor('#FFC7CE')
    
    # Add title
    plt.title(title, fontsize=14, fontweight='bold', pad=20)
    
    # Save with high resolution
    plt.savefig(output_path, dpi=150, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    plt.close()


def create_charts(orders_df: pd.DataFrame, summary_df: pd.DataFrame) -> dict:
    """
    Create visualization charts from the procurement data.
    
    Args:
        orders_df: Orders DataFrame.
        summary_df: Summary DataFrame.
        
    Returns:
        Dictionary with paths to all chart images.
    """
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    
    # Set style for professional look
    plt.style.use('seaborn-v0_8-whitegrid')
    
    # Color palette matching our green theme
    primary_color = '#217346'
    secondary_colors = ['#217346', '#2E9B5F', '#45B778', '#6FCF97', '#A8E6CF']
    
    charts = {}
    
    # Chart 1: Top 5 NDC by Spend (Horizontal Bar Chart)
    fig, ax = plt.subplots(figsize=(10, 6))
    # Group by NDC and Drug name for better labels
    ndc_spend = orders_df.groupby(['NDC', 'Drug'])['Order Value'].sum().reset_index()
    ndc_spend['Label'] = ndc_spend['Drug'] + '\n(' + ndc_spend['NDC'] + ')'
    top_5_ndc = ndc_spend.nlargest(5, 'Order Value').sort_values('Order Value', ascending=True)
    
    bars = ax.barh(top_5_ndc['Label'], top_5_ndc['Order Value'], color=secondary_colors[::-1])
    ax.set_xlabel('Total Spend ($)', fontsize=11)
    ax.set_title('Top 5 NDC by Spend', fontsize=14, fontweight='bold', pad=15)
    ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
    
    # Add value labels on bars
    for bar, value in zip(bars, top_5_ndc['Order Value']):
        ax.text(value + max(top_5_ndc['Order Value']) * 0.01, bar.get_y() + bar.get_height()/2, 
                f'${value:,.0f}', va='center', fontsize=10, fontweight='bold')
    
    plt.tight_layout()
    plt.savefig(str(CHART_TOP_NDC_SPEND), dpi=150, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    plt.close()
    charts['top_ndc_spend'] = str(CHART_TOP_NDC_SPEND)
    print(f"✓ Chart created: Top 5 NDC Spend")
    
    # Chart 2: Monthly Savings (Bar Chart)
    fig, ax = plt.subplots(figsize=(10, 6))
    summary_sorted = summary_df.sort_values('Month')
    
    # Color bars based on positive (green) or negative (red) savings
    colors = [primary_color if s >= 0 else '#DC3545' for s in summary_sorted['Savings $']]
    bars = ax.bar(summary_sorted['Month'], summary_sorted['Savings $'], color=colors)
    
    ax.set_xlabel('Month', fontsize=11)
    ax.set_ylabel('Savings ($)', fontsize=11)
    ax.set_title('Monthly Savings vs Previous Period', fontsize=14, fontweight='bold', pad=15)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
    ax.axhline(y=0, color='gray', linestyle='-', linewidth=0.5)
    plt.xticks(rotation=45, ha='right')
    
    # Add value labels on bars
    for bar, pct in zip(bars, summary_sorted['Savings %']):
        height = bar.get_height()
        label = f'${height:,.0f}\n({pct:+.1f}%)'
        va = 'bottom' if height >= 0 else 'top'
        ax.text(bar.get_x() + bar.get_width()/2., height, label, 
                ha='center', va=va, fontsize=9, fontweight='bold')
    
    plt.tight_layout()
    plt.savefig(str(CHART_SAVINGS_BY_MONTH), dpi=150, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    plt.close()
    charts['savings_by_month'] = str(CHART_SAVINGS_BY_MONTH)
    print(f"✓ Chart created: Monthly Savings")
    
    # Chart 3: Top 5 NDC Spend Distribution (Pie Chart)
    fig, ax = plt.subplots(figsize=(10, 8))
    ndc_spend = orders_df.groupby(['NDC', 'Drug'])['Order Value'].sum().reset_index()
    top_5 = ndc_spend.nlargest(5, 'Order Value')
    
    # Create labels with NDC
    labels = [f"{row['Drug']}\n({row['NDC']})" for _, row in top_5.iterrows()]
    
    wedges, texts, autotexts = ax.pie(
        top_5['Order Value'], 
        labels=labels,
        autopct=lambda pct: f'${pct/100*sum(top_5["Order Value"]):,.0f}\n({pct:.1f}%)',
        colors=secondary_colors,
        explode=[0.02] * 5,
        startangle=90,
        textprops={'fontsize': 9}
    )
    
    # Style the percentage labels
    for autotext in autotexts:
        autotext.set_fontweight('bold')
        autotext.set_fontsize(9)
    
    ax.set_title('Top 5 NDC Spend Distribution', fontsize=14, fontweight='bold', pad=15)
    
    plt.tight_layout()
    plt.savefig(str(CHART_TOP_NDC_PIE), dpi=150, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    plt.close()
    charts['top_ndc_pie'] = str(CHART_TOP_NDC_PIE)
    print(f"✓ Chart created: Top 5 NDC Pie Chart")
    
    return charts


def create_word_document(summary_image: str, orders_image: str, charts: dict = None) -> str:
    """
    Create a Word document with the report.
    
    Args:
        summary_image: Path to summary screenshot.
        orders_image: Path to orders screenshot.
        charts: Dictionary of chart image paths.
    Returns:
        Path to the created Word document.
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading("Pharmacy Procurement Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    timestamp_para = doc.add_paragraph()
    timestamp_run = timestamp_para.add_run(f"Generated: {timestamp}")
    timestamp_run.italic = True
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # Summary Section
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Monthly overview of pharmacy procurement activities, "
        "including total orders, spending, and savings indicators."
    )
    
    # Insert summary image
    if os.path.exists(summary_image):
        doc.add_picture(summary_image, width=Inches(6.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Summary image not found]")
    
    doc.add_paragraph()  # Spacer
    
    # Orders Section
    doc.add_heading("Orders", level=1)
    doc.add_paragraph(
        "Detailed order records showing hospital, pharmacy, drug information, "
        "pricing, and invoicing status."
    )
    
    # Insert orders image
    if os.path.exists(orders_image):
        doc.add_picture(orders_image, width=Inches(6.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Orders image not found]")
    
    # Visualizations Section
    if charts:
        doc.add_page_break()
        doc.add_heading("Visualizations & Analytics", level=1)
        doc.add_paragraph(
            "The following charts provide visual insights into NDC spending patterns "
            "and savings trends across the reporting period."
        )
        doc.add_paragraph()
        
        # Top 5 NDC Spend Bar Chart
        if 'top_ndc_spend' in charts and os.path.exists(charts['top_ndc_spend']):
            doc.add_heading("Top 5 NDC by Spend", level=2)
            doc.add_paragraph(
                "This chart shows the top 5 National Drug Codes (NDC) ranked by total spend, "
                "helping identify high-cost medications for contract negotiations and formulary review."
            )
            doc.add_picture(charts['top_ndc_spend'], width=Inches(6))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # Monthly Savings Chart
        if 'savings_by_month' in charts and os.path.exists(charts['savings_by_month']):
            doc.add_heading("Monthly Savings", level=2)
            doc.add_paragraph(
                "This chart displays the savings achieved each month compared to the previous period. "
                "Green bars indicate savings, while red bars indicate cost increases."
            )
            doc.add_picture(charts['savings_by_month'], width=Inches(6))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # Top 5 NDC Pie Chart
        if 'top_ndc_pie' in charts and os.path.exists(charts['top_ndc_pie']):
            doc.add_heading("Top 5 NDC Spend Distribution", level=2)
            doc.add_paragraph(
                "Pie chart showing the distribution of spend across the top 5 NDC codes, "
                "illustrating the concentration of pharmacy procurement costs."
            )
            doc.add_picture(charts['top_ndc_pie'], width=Inches(5.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
    
    # Add footer note
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "This report was automatically generated. "
        "For questions, contact the Procurement Department."
    )
    footer_run.font.size = Pt(9)
    footer_run.italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    word_path = str(WORD_FILE)
    doc.save(word_path)
    
    print(f"✓ Word document created: {word_path}")
    return word_path


def generate_full_report() -> dict:
    """
    Main function to generate the complete report package.
    
    Returns:
        Dictionary with paths to all generated files.
    """
    print("=" * 60)
    print("PHARMACY PROCUREMENT REPORT GENERATOR")
    print("=" * 60)
    print()
    
    # Ensure output directory exists
    ensure_output_directory()
    
    # Step 1: Get data from data source
    print("Step 1: Loading data...")
    summary_df, orders_df = get_report_data()
    print(f"  - Summary records: {len(summary_df)}")
    print(f"  - Order records: {len(orders_df)}")
    print()
    
    # Step 2: Create Excel report
    print("Step 2: Creating Excel report...")
    excel_path = create_excel_report(summary_df, orders_df)
    print()
    
    # Step 3: Create screenshots
    print("Step 3: Creating Excel screenshots...")
    summary_image, orders_image = create_excel_screenshots(excel_path)
    print()
    
    # Step 4: Create visualization charts
    print("Step 4: Creating visualization charts...")
    charts = create_charts(orders_df, summary_df)
    print()
    
    # Step 5: Create Word document
    print("Step 5: Creating Word document...")
    word_path = create_word_document(summary_image, orders_image, charts)
    print()
    
    # Summary
    print("=" * 60)
    print("REPORT GENERATION COMPLETE")
    print("=" * 60)
    print(f"Output directory: {OUTPUT_DIR}")
    print()
    print("Generated files:")
    print(f"  • {EXCEL_FILE.name}")
    print(f"  • {SUMMARY_IMAGE.name}")
    print(f"  • {ORDERS_IMAGE.name}")
    print(f"  • {CHART_TOP_NDC_SPEND.name}")
    print(f"  • {CHART_SAVINGS_BY_MONTH.name}")
    print(f"  • {CHART_TOP_NDC_PIE.name}")
    print(f"  • {WORD_FILE.name}")
    
    return {
        "excel": excel_path,
        "summary_image": summary_image,
        "orders_image": orders_image,
        "charts": charts,
        "word": word_path,
    }


# Main entry point
if __name__ == "__main__":
    generate_full_report()
